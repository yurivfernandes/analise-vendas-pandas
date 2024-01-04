import os
from docx import Document

def listar_nomes_arquivos(caminho_pasta):
    nomes_arquivos = []

    for arquivo in os.listdir(caminho_pasta):
        caminho_arquivo = os.path.join(caminho_pasta, arquivo)
        if os.path.isfile(caminho_arquivo):
            nomes_arquivos.append(arquivo)

    return nomes_arquivos

def salvar_em_word(lista_nomes_arquivos, caminho_pasta, caminho_salvar):
    nome_arquivo_word = os.path.join(caminho_salvar, 'lista_arquivos.docx')
    document = Document()
    document.add_heading(f'Lista de Arquivos enviados pelo cliente: ', level=1)

    for nome_arquivo in lista_nomes_arquivos:
        document.add_paragraph(nome_arquivo)

    document.save(nome_arquivo_word)
    print(f'Lista de arquivos salva em {nome_arquivo_word}')

caminho_do_diretorio = input("Digite o caminho da pasta: ")
caminho_salvar = input("Digite o caminho onde deseja salvar o arquivo com a lista: ")

lista_nomes_arquivos = listar_nomes_arquivos(caminho_do_diretorio)
salvar_em_word(lista_nomes_arquivos, caminho_do_diretorio, caminho_salvar)

input("Pressione Enter para fechar.")